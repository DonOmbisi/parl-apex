import logging
from datetime import datetime
from pathlib import Path

from docx import Document

logger = logging.getLogger("parl_apex.documents.quotation_generator")

TEMPLATE_PATH = Path("documents/templates/quotation_template.docx")
OUTPUT_DIR = Path("data/quotations")


def format_ordinal_date(date: datetime) -> str:
    """Format date as ordinal (e.g., '12th June 2026')."""
    day = date.day
    if 11 <= day <= 13:
        suffix = "th"
    else:
        suffix = {1: "st", 2: "nd", 3: "rd"}.get(day % 10, "th")
    return date.strftime(f"{day}{suffix} %B %Y")


def format_number(amount: float) -> str:
    """Format number with commas and 2 decimal places."""
    return f"{amount:,.2f}"


def replace_placeholders(doc: Document, data: dict) -> Document:
    """
    Replace placeholders in the document with actual values using run-level text replacement.
    
    Iterates through all paragraphs and table cells' runs to find and replace tokens.
    """
    # Extract data fields
    client_company = data.get("client_company", "")
    contact_person = data.get("contact_person", "")
    client_address = data.get("client_address", "")
    requirement = data.get("requirement", "")
    salesperson = data.get("salesperson", "")
    payment_terms = data.get("payment_terms", "")
    valid_until = data.get("valid_until", "")
    quotation_number = data.get("quotation_number", "")
    delivery_method = data.get("delivery_method", "")
    line_item_name = data.get("line_item_name", "")
    total_cost = data.get("total_cost", 0)
    sub_total = data.get("sub_total", 0)
    discount_amount = data.get("discount_amount", 0)
    
    # Format dates
    today = datetime.now()
    today_formatted = format_ordinal_date(today)
    
    # Build replacement mapping - exact strings from template
    replacements = {
        "World Food Program WFP – Kenya/RBN Office": client_company,
        "Zaccheus Ndirima - zaccheus.ndirima@wfp.org": contact_person,
        "UN Gigiri Compound, P.O. Box 44482, Nairobi, Kenya": client_address,
        "PS IMAGO PRO (An IBM SPSS Statistics Solution)": requirement,
        "Mary Nthambi": salesperson,
        "30 Days after issuance of LPO.": payment_terms,
        "12th June 2026": valid_until,
        "P001 1205 2026": quotation_number,
        "Download Link": delivery_method,
        "{{ line_item_name }}": line_item_name,
        "{{ total }}": format_number(total_cost),
        "12,645.00": format_number(sub_total),
        "2,529.00": format_number(discount_amount),
    }
    
    # Also replace the date in the title
    # Pattern: "SOFTWARE QUOTATION {date}"
    title_date_replacements = {
        "12th May 2026": today_formatted,
    }
    
    # Helper function to replace text in runs
    def replace_in_paragraph(paragraph, mapping):
        for run in paragraph.runs:
            for old_text, new_value in mapping.items():
                if old_text in run.text:
                    run.text = run.text.replace(old_text, str(new_value))
    
    # Replace in paragraphs
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph, replacements)
        replace_in_paragraph(paragraph, title_date_replacements)
    
    # Replace in table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph, replacements)
                    replace_in_paragraph(paragraph, title_date_replacements)
    
    return doc


def generate_quotation(data: dict) -> str:
    """
    Generate a filled DOCX quotation from the template.
    
    Args:
        data: Dictionary containing quotation fields
    
    Returns:
        Path to the generated quotation file
    """
    # Ensure output directory exists
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    
    # Check if template exists
    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"Quotation template not found at {TEMPLATE_PATH}")
    
    # Load the template
    doc = Document(TEMPLATE_PATH)
    
    # Replace placeholders with actual data
    doc = replace_placeholders(doc, data)
    
    # Get quotation number from data
    quotation_number = data.get("quotation_number", "unknown")
    output_path = OUTPUT_DIR / f"{quotation_number}.docx"
    
    # Save the document
    doc.save(output_path)
    logger.info(f"Generated quotation at {output_path}")
    
    return str(output_path)
